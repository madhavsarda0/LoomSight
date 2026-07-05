#!/usr/bin/env python3
"""
WeaveCAD — Full-stack web app for converting greige fabric technical sheets
into weave pattern CAD images (PNG/SVG).

Run:  python app.py
Open: http://localhost:5000
"""

import os
import re
import io
import base64
import json
import tempfile
from datetime import datetime
from pathlib import Path

from flask import Flask, render_template, request, jsonify, send_file, Response
from werkzeug.utils import secure_filename
from PIL import Image, ImageDraw

# Optional imports — graceful degradation if not installed
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import pytesseract
    from PIL import Image as PILImage
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False


app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32 MB
app.config['UPLOAD_FOLDER'] = 'uploads'

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'pdf', 'jpg', 'jpeg', 'png', 'doc', 'docx', 'xls', 'xlsx', 'csv', 'txt'}


# ───────────────────────────────────────────────────────────────
# Weave pattern definitions
# ───────────────────────────────────────────────────────────────

WEAVE_MATRICES = {
    'plain':    [[1,0],[0,1]],
    'twill12':  [[1,0,0],[0,1,0],[0,0,1]],
    'twill21':  [[0,1,1],[1,0,1],[1,1,0]],
    'twill22':  [[1,1,0,0],[0,1,1,0],[0,0,1,1],[1,0,0,1]],
    'twill31':  [[1,0,0,0],[0,1,0,0],[0,0,1,0],[0,0,0,1]],
    'twill13':  [[0,1,1,1],[1,0,1,1],[1,1,0,1],[1,1,1,0]],
    'satin5':   [[0,0,1,0,0],[0,0,0,0,1],[1,0,0,0,0],[0,0,0,1,0],[0,1,0,0,0]],
    'satin8':   [[0,0,0,1,0,0,0,0],[0,0,0,0,0,0,1,0],[0,1,0,0,0,0,0,0],[0,0,0,0,1,0,0,0],
                  [0,0,0,0,0,0,0,1],[0,0,1,0,0,0,0,0],[1,0,0,0,0,0,0,0],[0,0,0,0,0,1,0,0]],
    'basket22': [[1,1,0,0],[1,1,0,0],[0,0,1,1],[0,0,1,1]],
    'basket33': [[1,1,1,0,0,0],[1,1,1,0,0,0],[1,1,1,0,0,0],[0,0,0,1,1,1],[0,0,0,1,1,1],[0,0,0,1,1,1]],
    'huck':     [[1,1,0,0,1,1],[1,1,0,0,1,1],[0,0,1,1,0,0],[0,0,1,1,0,0],[1,1,0,0,1,1],[1,1,0,0,1,1]],
}

WEAVE_NAMES = {
    'plain': 'Plain 1/1',
    'twill12': 'Twill 1/2',
    'twill21': 'Twill 2/1',
    'twill22': 'Twill 2/2',
    'twill31': 'Twill 3/1',
    'twill13': 'Twill 1/3',
    'satin5': 'Satin 5-harness',
    'satin8': 'Satin 8-harness',
    'basket22': 'Basket 2/2',
    'basket33': 'Basket 3/3',
    'huck': 'Huckaback',
}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ───────────────────────────────────────────────────────────────
# Text extraction
# ───────────────────────────────────────────────────────────────

def extract_text_from_pdf(filepath):
    """Extract text from PDF. Falls back to OCR if pdfplumber finds nothing."""
    text = ""
    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            text += f"[pdfplumber error: {e}]\n"

    # If no text found, try OCR on rendered pages
    if not text.strip() and HAS_TESSERACT:
        try:
            # Convert PDF pages to images using pdf2image if available
            try:
                from pdf2image import convert_from_path
                images = convert_from_path(filepath, dpi=200)
                for img in images:
                    text += pytesseract.image_to_string(img) + "\n"
            except ImportError:
                text += "[PDF appears to be scanned. Install pdf2image + poppler for OCR support.]\n"
        except Exception as e:
            text += f"[OCR error: {e}]\n"

    return text


def extract_text_from_image(filepath):
    """OCR on image files."""
    if not HAS_TESSERACT:
        return "[Tesseract OCR not installed. Install pytesseract and tesseract-ocr.]"
    try:
        img = PILImage.open(filepath)
        # Convert to RGB if necessary
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')
        return pytesseract.image_to_string(img)
    except Exception as e:
        return f"[Image OCR error: {e}]"


def extract_text_from_excel(filepath):
    """Extract text from Excel files."""
    text = ""
    if HAS_PANDAS:
        try:
            # Try all sheets
            xls = pd.ExcelFile(filepath)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                text += f"\n--- Sheet: {sheet_name} ---\n"
                text += df.to_string(index=False, header=True) + "\n"
        except Exception as e:
            text += f"[pandas error: {e}]\n"
    elif HAS_OPENPYXL:
        try:
            wb = openpyxl.load_workbook(filepath, data_only=True)
            for sheet in wb.worksheets:
                text += f"\n--- Sheet: {sheet.title} ---\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        text += row_text + "\n"
        except Exception as e:
            text += f"[openpyxl error: {e}]\n"
    else:
        text = "[Install openpyxl or pandas to read Excel files.]"
    return text


def extract_text_from_docx(filepath):
    """Extract text from Word documents."""
    if not HAS_DOCX:
        return "[python-docx not installed. Install python-docx to read DOCX files.]"
    try:
        doc = Document(filepath)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    text += "\n" + row_text
        return text
    except Exception as e:
        return f"[DOCX error: {e}]"


def extract_text_from_doc(filepath):
    """Extract text from old .doc files."""
    try:
        # Try antiword or textract approaches
        import subprocess
        result = subprocess.run(['antiword', filepath], capture_output=True, text=True)
        if result.returncode == 0:
            return result.stdout
        else:
            return f"[antiword failed: {result.stderr}. Convert .doc to .docx first.]"
    except FileNotFoundError:
        return "[antiword not found. Install antiword or convert .doc to .docx.]"
    except Exception as e:
        return f"[DOC extraction error: {e}]"


def extract_text(filepath, ext):
    """Route to appropriate extractor based on file extension."""
    ext = ext.lower()
    if ext == 'pdf':
        return extract_text_from_pdf(filepath)
    elif ext in ('jpg', 'jpeg', 'png'):
        return extract_text_from_image(filepath)
    elif ext in ('xls', 'xlsx'):
        return extract_text_from_excel(filepath)
    elif ext == 'docx':
        return extract_text_from_docx(filepath)
    elif ext == 'doc':
        return extract_text_from_doc(filepath)
    elif ext in ('txt', 'csv'):
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    return "[Unsupported file type]"


# ───────────────────────────────────────────────────────────────
# Text parsing for weave specs
# ───────────────────────────────────────────────────────────────

def parse_specs(text):
    """Parse extracted text for weaving parameters."""
    t = text.lower()
    specs = {
        'epi': None,
        'ppi': None,
        'warp_count': None,
        'weft_count': None,
        'width': None,
        'weave': None,
        'construction': None,
    }

    # EPI / ends per inch
    epi_match = re.search(r'(?:epi|ends?\s*per\s*inch|warp\s*density|ends)[\s:]*(\d+)', t, re.IGNORECASE)
    if epi_match:
        specs['epi'] = int(epi_match.group(1))

    # PPI / picks per inch
    ppi_match = re.search(r'(?:ppi|picks?\s*per\s*inch|weft\s*density|picks)[\s:]*(\d+)', t, re.IGNORECASE)
    if ppi_match:
        specs['ppi'] = int(ppi_match.group(1))

    # Width
    width_match = re.search(r'(?:width|reed\s*width|fabric\s*width)[\s:]*(\d+(?:\.\d+)?)', t, re.IGNORECASE)
    if width_match:
        specs['width'] = float(width_match.group(1))

    # Warp / Weft count
    warp_match = re.search(r'(?:warp\s*count|warp\s*yarn|warp)[\s:]*(\d+\/?\d*)', t, re.IGNORECASE)
    if warp_match:
        specs['warp_count'] = warp_match.group(1)

    weft_match = re.search(r'(?:weft\s*count|weft\s*yarn|filling\s*count|weft)[\s:]*(\d+\/?\d*)', t, re.IGNORECASE)
    if weft_match:
        specs['weft_count'] = weft_match.group(1)

    # Construction format: 80x72/40x40 or 80 X 72 / 40 X 40
    construction = re.search(r'(\d+)[xX](\d+)\/(\d+\/?\d*)[xX](\d+\/?\d*)', text)
    if construction:
        specs['epi'] = int(construction.group(1))
        specs['ppi'] = int(construction.group(2))
        specs['warp_count'] = construction.group(3)
        specs['weft_count'] = construction.group(4)
        specs['construction'] = f"{construction.group(1)}x{construction.group(2)}/{construction.group(3)}x{construction.group(4)}"

    # Weave type detection
    weave_map = {
        'plain': re.compile(r'plain|1[-/]1', re.IGNORECASE),
        'twill12': re.compile(r'twill\s*1[-/]2|1[-/]2\s*twill', re.IGNORECASE),
        'twill21': re.compile(r'twill\s*2[-/]1|2[-/]1\s*twill', re.IGNORECASE),
        'twill22': re.compile(r'twill\s*2[-/]2|2[-/]2\s*twill', re.IGNORECASE),
        'twill31': re.compile(r'twill\s*3[-/]1|3[-/]1\s*twill', re.IGNORECASE),
        'twill13': re.compile(r'twill\s*1[-/]3|1[-/]3\s*twill', re.IGNORECASE),
        'satin5': re.compile(r'satin\s*5|5[-\s]?harness', re.IGNORECASE),
        'satin8': re.compile(r'satin\s*8|8[-\s]?harness', re.IGNORECASE),
        'basket22': re.compile(r'basket\s*2[-/]2|2[-/]2\s*basket', re.IGNORECASE),
        'basket33': re.compile(r'basket\s*3[-/]3|3[-/]3\s*basket', re.IGNORECASE),
        'huck': re.compile(r'huckaback|huck', re.IGNORECASE),
    }

    for key, pattern in weave_map.items():
        if pattern.search(text):
            specs['weave'] = key
            break

    return specs


# ───────────────────────────────────────────────────────────────
# Weave pattern generation
# ───────────────────────────────────────────────────────────────

def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def shade_color(hex_color, percent):
    r, g, b = hex_to_rgb(hex_color)
    R = min(255, int(r * (1 + percent / 100)))
    G = min(255, int(g * (1 + percent / 100)))
    B = min(255, int(b * (1 + percent / 100)))
    return (R, G, B)


def generate_weave_png(weave_type, epi, ppi, warp_color, weft_color, repeats, view_mode, size=800):
    """Generate weave pattern as PNG image."""
    matrix = WEAVE_MATRICES.get(weave_type, WEAVE_MATRICES['plain'])
    h = len(matrix)
    w = len(matrix[0]) if matrix else 1
    rows = h * repeats
    cols = w * repeats

    img = Image.new('RGB', (size, size), hex_to_rgb(weft_color))
    draw = ImageDraw.Draw(img)

    cell_w = size / cols
    cell_h = size / rows

    if view_mode == 'grid':
        # Point paper view
        for i in range(rows):
            for j in range(cols):
                is_warp = matrix[i % h][j % w]
                color = hex_to_rgb(warp_color) if is_warp else hex_to_rgb(weft_color)
                x0 = j * cell_w
                y0 = i * cell_h
                x1 = (j + 1) * cell_w - 0.5
                y1 = (i + 1) * cell_h - 0.5
                draw.rectangle([x0, y0, x1, y1], fill=color, outline=(220, 220, 220))
    else:
        # Fabric simulation view
        # Draw weft threads first (background layer)
        for i in range(rows):
            for j in range(cols):
                is_warp = matrix[i % h][j % w]
                if not is_warp:
                    # Weft thread visible
                    x0 = j * cell_w
                    y0 = i * cell_h + cell_h * 0.15
                    x1 = (j + 1) * cell_w
                    y1 = i * cell_h + cell_h * 0.85
                    draw.rectangle([x0, y0, x1, y1], fill=shade_color(weft_color, -8))
                    # Shadow
                    draw.rectangle([x0, y1, x1, i * cell_h + cell_h], fill=(0, 0, 0, 15))

        # Draw warp threads on top
        for i in range(rows):
            for j in range(cols):
                is_warp = matrix[i % h][j % w]
                if is_warp:
                    # Warp thread visible (over weft)
                    x0 = j * cell_w + cell_w * 0.15
                    y0 = i * cell_h
                    x1 = j * cell_w + cell_w * 0.85
                    y1 = (i + 1) * cell_h
                    draw.rectangle([x0, y0, x1, y1], fill=hex_to_rgb(warp_color))
                    # Highlight on top
                    draw.rectangle([x0, y0, x1, y0 + cell_h * 0.3], fill=shade_color(warp_color, 12))
                    # Shadow on right
                    draw.rectangle([x1, y0, j * cell_w + cell_w, y1], fill=(0, 0, 0, 20))

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def generate_weave_svg(weave_type, epi, ppi, warp_color, weft_color, repeats, view_mode, size=800):
    """Generate weave pattern as SVG."""
    matrix = WEAVE_MATRICES.get(weave_type, WEAVE_MATRICES['plain'])
    h = len(matrix)
    w = len(matrix[0]) if matrix else 1
    rows = h * repeats
    cols = w * repeats

    cell_w = size / cols
    cell_h = size / rows

    svg_parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 {size} {size}">',
        f'<rect width="{size}" height="{size}" fill="{weft_color}"/>'
    ]

    if view_mode == 'grid':
        for i in range(rows):
            for j in range(cols):
                is_warp = matrix[i % h][j % w]
                fill = warp_color if is_warp else weft_color
                x = j * cell_w
                y = i * cell_h
                cw = cell_w - 0.5
                ch = cell_h - 0.5
                svg_parts.append(
                    f'<rect x="{x:.2f}" y="{y:.2f}" width="{cw:.2f}" height="{ch:.2f}" '
                    f'fill="{fill}" stroke="rgba(0,0,0,0.08)" stroke-width="0.5"/>'
                )
    else:
        # Fabric simulation SVG
        for i in range(rows):
            for j in range(cols):
                is_warp = matrix[i % h][j % w]
                if not is_warp:
                    x = j * cell_w
                    y = i * cell_h + cell_h * 0.15
                    svg_parts.append(
                        f'<rect x="{x:.2f}" y="{y:.2f}" width="{cell_w:.2f}" height="{cell_h*0.7:.2f}" fill="{weft_color}"/>'
                    )
                    svg_parts.append(
                        f'<rect x="{x:.2f}" y="{y + cell_h*0.7:.2f}" width="{cell_w:.2f}" height="{cell_h*0.15:.2f}" fill="rgba(0,0,0,0.06)"/>'
                    )

        for i in range(rows):
            for j in range(cols):
                is_warp = matrix[i % h][j % w]
                if is_warp:
                    x = j * cell_w + cell_w * 0.15
                    y = i * cell_h
                    svg_parts.append(
                        f'<rect x="{x:.2f}" y="{y:.2f}" width="{cell_w*0.7:.2f}" height="{cell_h:.2f}" fill="{warp_color}"/>'
                    )
                    svg_parts.append(
                        f'<rect x="{x:.2f}" y="{y:.2f}" width="{cell_w*0.7:.2f}" height="{cell_h*0.3:.2f}" fill="{warp_color}" opacity="0.7"/>'
                    )
                    svg_parts.append(
                        f'<rect x="{x + cell_w*0.7:.2f}" y="{y:.2f}" width="{cell_w*0.15:.2f}" height="{cell_h:.2f}" fill="rgba(0,0,0,0.08)"/>'
                    )

    svg_parts.append('</svg>')
    return '\n'.join(svg_parts)


# ───────────────────────────────────────────────────────────────
# Flask routes
# ───────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        # Add timestamp to avoid collisions
        name, ext = os.path.splitext(filename)
        filename = f"{name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Extract text
        ext_clean = filename.rsplit('.', 1)[1].lower()
        raw_text = extract_text(filepath, ext_clean)

        # Parse specs
        specs = parse_specs(raw_text)

        return jsonify({
            'success': True,
            'filename': filename,
            'raw_text': raw_text,
            'specs': specs
        })

    return jsonify({'error': 'File type not allowed'}), 400


@app.route('/api/parse', methods=['POST'])
def parse_text():
    data = request.get_json()
    text = data.get('text', '')
    specs = parse_specs(text)
    return jsonify({'specs': specs})


@app.route('/api/generate', methods=['POST'])
def generate_pattern():
    data = request.get_json()

    weave_type = data.get('weave', 'plain')
    epi = int(data.get('epi', 80))
    ppi = int(data.get('ppi', 72))
    warp_color = data.get('warp_color', '#e0e0e0')
    weft_color = data.get('weft_color', '#b0b0b0')
    repeats = int(data.get('repeats', 10))
    view_mode = data.get('view_mode', 'grid')
    size = int(data.get('size', 800))

    # Generate PNG
    png_buf = generate_weave_png(weave_type, epi, ppi, warp_color, weft_color, repeats, view_mode, size)
    png_b64 = base64.b64encode(png_buf.read()).decode('utf-8')

    # Generate SVG
    svg_content = generate_weave_svg(weave_type, epi, ppi, warp_color, weft_color, repeats, view_mode, size)
    svg_b64 = base64.b64encode(svg_content.encode('utf-8')).decode('utf-8')

    # Calculate summary
    matrix = WEAVE_MATRICES.get(weave_type, WEAVE_MATRICES['plain'])
    repeat_h = len(matrix)
    repeat_w = len(matrix[0]) if matrix else 1
    width = data.get('width')
    total_ends = round(epi * float(width)) if width else None

    return jsonify({
        'success': True,
        'png': png_b64,
        'svg': svg_b64,
        'summary': {
            'weave_name': WEAVE_NAMES.get(weave_type, 'Unknown'),
            'construction': f"{epi}x{ppi}",
            'width': f"{width} in" if width else '—',
            'total_ends': total_ends if total_ends else '—',
            'repeat_size': f"{repeat_w} × {repeat_h}",
            'view_mode': 'Point paper' if view_mode == 'grid' else 'Fabric simulation'
        }
    })


@app.route('/api/download/<format>', methods=['POST'])
def download_file(format):
    data = request.get_json()

    weave_type = data.get('weave', 'plain')
    epi = int(data.get('epi', 80))
    ppi = int(data.get('ppi', 72))
    warp_color = data.get('warp_color', '#e0e0e0')
    weft_color = data.get('weft_color', '#b0b0b0')
    repeats = int(data.get('repeats', 10))
    view_mode = data.get('view_mode', 'grid')
    size = int(data.get('size', 800))

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    if format == 'png':
        png_buf = generate_weave_png(weave_type, epi, ppi, warp_color, weft_color, repeats, view_mode, size)
        return send_file(
            png_buf,
            mimetype='image/png',
            as_attachment=True,
            download_name=f'weave_{weave_type}_{timestamp}.png'
        )
    elif format == 'svg':
        svg_content = generate_weave_svg(weave_type, epi, ppi, warp_color, weft_color, repeats, view_mode, size)
        svg_buf = io.BytesIO(svg_content.encode('utf-8'))
        return send_file(
            svg_buf,
            mimetype='image/svg+xml',
            as_attachment=True,
            download_name=f'weave_{weave_type}_{timestamp}.svg'
        )

    return jsonify({'error': 'Invalid format'}), 400


import os

# ───────────────────────────────────────────────────────────────
# Health check endpoint (for Render monitoring)
# ───────────────────────────────────────────────────────────────

@app.route('/api/health')
def health_check():
    return jsonify({'status': 'ok', 'service': 'weave-cad', 'timestamp': datetime.now().isoformat()})

# ───────────────────────────────────────────────────────────────
# Main
# ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)

